import pandas as pd
import locale
from docxcompose.composer import Composer
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
from docx.shared import Inches, Pt

def change_font(doc, font_name, font_size, top_margin, bottom_margin, left_margin, right_margin):
    # Cambia el tipo y tamaño de letra de todos los párrafos en el documento
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # Cambia el tipo y tamaño de letra de todas las tablas en el documento
    for table in doc.tables:
        # Configura la tabla para que ajuste el ancho automáticamente
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size-1)

    # Cambiar márgenes de las secciones del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top_margin)
        section.bottom_margin = Inches(bottom_margin)
        section.left_margin = Inches(left_margin)
        section.right_margin = Inches(right_margin)
    
    # Retornar el documento modificado
    return doc